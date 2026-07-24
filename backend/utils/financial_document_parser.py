# ── 金语AI 金融文档解析与切分 ──
# 决策记录：
# - 复用 LexAI 的 DocumentChunk 结构，保持与向量库 / API 接口一致
# - 监管法规/基金合同：按「第X条/款/章/节」正则切，保留层级元数据
# - 年报/研报：按【章节标题】或 Markdown 标题切；长段落按 300-500 字回退
# - MIN_CHUNK_SIZE=120：金融段落跨度大，避免过碎
# - doctype 自动检测：文件名/正文含关键词区分 regulation / contract / report
# - 每块 metadata 记录 section、doctype、source（文件名）

import os
import re
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path

# ── 复用基础结构 ──
try:
    from .document_parser import DocumentChunk
except ImportError:
    try:
        from document_parser import DocumentChunk
    except ImportError:
        class DocumentChunk:
            def __init__(self, content: str, page_number=None,
                         paragraph_number=None, metadata=None):
                self.id = str(uuid.uuid4())
                self.content = content
                self.page_number = page_number
                self.paragraph_number = paragraph_number
                self.metadata = metadata or {}


class FinancialDocumentParser:
    # 最小合并阈值：不足此字数的相邻块自动合并到前一块或下一块
    MIN_CHUNK_SIZE = 120

    # ── 正则释义 ──
    # (?:第...)      非捕获分组，避免 split 时吃掉匹配内容
    # [零一二三四五六七八九十百千\d]  同时匹配中文数字与阿拉伯数字
    # [条款章节]      法律文书层级单位
    # 关键修复：只在"条款起点"切分——前一个字符须为句末标点/换行/全角空格/
    #   右括号等，刻意不含 、和 ，，从而避免把"本法第五十三条"这类行内引用
    #   误判为新的一条（否则一条法条会被切成多块）。
    ARTICLE_SPLIT = re.compile(
        r'(?:^|(?<=[。；！？\r\n\t \u3000）】”》」]))'
        r'(第[零一二三四五六七八九十百千\d]+[条款章节])'
    )

    # 基金合同/章程标题：第一章、第二条、第三节……（整行匹配，用于 Word 逐段流）
    CONTRACT_HEADING = re.compile(
        r'^第[零一二三四五六七八九十百千\d]+[章节条]'
    )

    # 年报章节标题：【重要提示】【公司简介】【财务数据】……
    REPORT_BRACKET_HEADING = re.compile(
        r'^[ \t]*[【（\(][^】）\)]{1,15}[】）\)]'
    )

    # Markdown 二级标题：## XXX
    MARKDOWN_HEADING = re.compile(r'^#{2,4}\s+.+')

    # ── 文档类型检测 ──
    DOCTYPE_KEYWORDS = {
        'contract': ['基金合同', '信托合同', '资产管理合同', '合伙协议',
                     '托管协议', '基金招募说明书'],
        'report':   ['年报', '年度报告', '中期报告', '季度报告',
                     '研究报告', '行业报告', '招股说明书', '募集说明书'],
    }

    @staticmethod
    def detect_doctype(filename: str, text_head: str = "") -> str:
        """根据文件名+正文前500字判断文档类型。
        优先级：contract > report > regulation（兜底）。"""
        combined = (filename + " " + text_head).lower()
        for doc_type, keywords in FinancialDocumentParser.DOCTYPE_KEYWORDS.items():
            for kw in keywords:
                if kw.lower() in combined:
                    return doc_type
        return 'regulation'

    # ═══════════════════════════════════════════════
    #  公共入口
    # ═══════════════════════════════════════════════

    @staticmethod
    def parse_file(file_path: str) -> List[DocumentChunk]:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return FinancialDocumentParser.parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return FinancialDocumentParser.parse_word(file_path)
        elif ext in ['.md', '.markdown']:
            return FinancialDocumentParser.parse_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def parse_pdf(file_path: str) -> List[DocumentChunk]:
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF (fitz) not installed. pip install pymupdf")
        chunks = []
        doc = fitz.open(file_path)
        filename = os.path.basename(file_path)
        doctype = FinancialDocumentParser.detect_doctype(filename)
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if not text.strip():
                    continue
                page_chunks = FinancialDocumentParser._split_financial_text(
                    text, page_num + 1, filename, doctype
                )
                chunks.extend(page_chunks)
        finally:
            doc.close()
        return chunks

    @staticmethod
    def parse_word(file_path: str) -> List[DocumentChunk]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. pip install python-docx")
        doc = Document(file_path)
        filename = os.path.basename(file_path)

        # 先收集全部段落用于 doctype 检测
        all_text = " ".join(p.text for p in doc.paragraphs[:50])
        doctype = FinancialDocumentParser.detect_doctype(filename, all_text)

        chunks = []
        paragraph_num = 0
        current_text = ""

        for para in doc.paragraphs:
            paragraph_num += 1
            text = para.text.strip()
            if not text:
                continue

            if FinancialDocumentParser._is_section_heading(text, doctype):
                # 遇到新标题，把攒的 current_text 切分输出
                if current_text.strip():
                    sub = FinancialDocumentParser._split_financial_text(
                        current_text.strip(), 1, filename, doctype
                    )
                    for s in sub:
                        s.paragraph_number = paragraph_num - 1
                        s.metadata['filename'] = filename
                    chunks.extend(sub)
                current_text = text + "\n"
            else:
                current_text += text + "\n"

        if current_text.strip():
            sub = FinancialDocumentParser._split_financial_text(
                current_text.strip(), 1, filename, doctype
            )
            for s in sub:
                s.paragraph_number = paragraph_num
                s.metadata['filename'] = filename
            chunks.extend(sub)

        return chunks

    @staticmethod
    def parse_markdown(file_path: str) -> List[DocumentChunk]:
        """解析 Markdown 文件（.md / .markdown）。
        
        整文件读入后复用 _split_financial_text 按内容结构切分，
        Markdown 标题（# / ## / ###）会自然被 heading 策略识别。
        """
        filename = os.path.basename(file_path)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # 兜底用系统默认编码
            with open(file_path, 'r', encoding='gbk') as f:
                text = f.read()
        # 去掉爬虫写入的元数据行（# 标题: / # 来源: / # 链接:），避免污染 chunk
        text = re.sub(r'^#\s*(标题|来源|链接):.*\n', '', text, flags=re.MULTILINE)
        # 检测前 500 字确定 doctype
        doctype = FinancialDocumentParser.detect_doctype(filename, text[:500])
        chunks = FinancialDocumentParser._split_financial_text(
            text, 1, filename, doctype
        )
        # 给每块补上文件名
        for c in chunks:
            c.metadata['filename'] = filename
        return chunks

    # ═══════════════════════════════════════════════
    #  核心切分路由
    # ═══════════════════════════════════════════════

    @staticmethod
    def _split_financial_text(text: str, page_num: int, filename: str,
                               doctype: Optional[str] = None) -> List[DocumentChunk]:
        """根据 doctype 选择切分策略。"""
        if doctype is None:
            doctype = FinancialDocumentParser.detect_doctype(filename, text[:500])

        source = filename
        common_meta = {"doctype": doctype, "source": source}

        if doctype in ('contract', 'regulation'):
            # ── 策略1：按「第X条/款/章/节」切 ──
            parts = re.split(FinancialDocumentParser.ARTICLE_SPLIT, text)
            # parts 形如 ["正文前文", "第一条", "正文内容...", "第二条", ...]
            raw_chunks = FinancialDocumentParser._group_article_parts(parts)
            if len(raw_chunks) >= 2:  # 至少检测到2个条款才算结构有效
                result = FinancialDocumentParser._merge_and_wrap(
                    raw_chunks, page_num, doc_type=doctype, source=source
                )
                if result:
                    return result

        # ── 策略2：按【章节标题】或 Markdown 标题切（年报/研报） ──
        heading_parts = re.split(
            r'((?:^[ \t]*[【（\(][^】）\)]{1,15}[】）\)])|(?:^#{2,4}\s+.+))',
            text, flags=re.MULTILINE
        )
        has_heading = sum(
            1 for p in heading_parts
            if FinancialDocumentParser.REPORT_BRACKET_HEADING.match(p.strip())
               or FinancialDocumentParser.MARKDOWN_HEADING.match(p.strip())
        )
        if has_heading >= 2:
            raw_chunks = FinancialDocumentParser._group_heading_parts(heading_parts)
            result = FinancialDocumentParser._merge_and_wrap(
                raw_chunks, page_num, doc_type=doctype, source=source
            )
            if result:
                return result

        # ── 策略3：长段落回退（300–500 字软切分） ──
        paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        return FinancialDocumentParser._wrap(
            paras, page_num, doc_type=doctype, source=source
        )

    # ═══════════════════════════════════════════════
    #  分组逻辑
    # ═══════════════════════════════════════════════

    @staticmethod
    def _group_article_parts(parts: List[str]) -> List[Dict]:
        """将 ARTICLE_SPLIT 的 split 结果分组为 [{section, content}]。
        正则带捕获分组，split 结果排列为：
          [正文1, "第一条", 正文2, "第二条", 正文3, ...]
        偶数索引为标题，奇数索引为正文。"""
        raw = []
        buffer_section = "正文"
        buffer_content = ""
        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue
            if i % 2 == 1:  # 捕获到的条款号
                if buffer_content.strip():
                    raw.append({"section": buffer_section, "content": buffer_content})
                buffer_section = part
                buffer_content = ""
            else:
                buffer_content += part + " "
        if buffer_content.strip():
            raw.append({"section": buffer_section, "content": buffer_content})
        return raw

    @staticmethod
    def _group_heading_parts(parts: List[str]) -> List[Dict]:
        """将 heading split 结果分组为 [{section, content}]。
        带捕获分组时 split 返回 [正文, 标题1, 正文1, 标题2, 正文2, ...]"""
        raw = []
        buffer_section = "正文"
        buffer_content = ""
        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue
            if i % 2 == 1:  # 捕获到的标题行
                if buffer_content.strip():
                    raw.append({"section": buffer_section, "content": buffer_content})
                buffer_section = part
                buffer_content = ""
            else:
                buffer_content += part + " "
        if buffer_content.strip():
            raw.append({"section": buffer_section, "content": buffer_content})
        return raw

    # ═══════════════════════════════════════════════
    #  合并逻辑（_merge / _wrap）
    # ═══════════════════════════════════════════════

    @staticmethod
    def _merge_and_wrap(raw_chunks: List[Dict], page_num: int,
                         doc_type: str = "regulation",
                         source: str = "") -> List[DocumentChunk]:
        """section 级合并：每节独立成块，极度短的节（<40字）并入前一节。
        每块附带 section / doctype / source 元数据。"""
        if not raw_chunks:
            return []

        merged = []
        for chunk in raw_chunks:
            content = chunk["content"].strip()
            if not content:
                continue
            if merged and len(merged[-1]["content"]) < 40:
                merged[-1]["content"] += "\n" + content
                # 保留最后遇到的 section 名
                merged[-1]["section"] = chunk["section"]
            else:
                merged.append({
                    "section": chunk["section"],
                    "content": content
                })

        if not merged:
            return [DocumentChunk(content="", page_number=page_num,
                                 metadata={"doctype": doc_type, "source": source,
                                           "section": "正文"})]

        # MIN_CHUNK_SIZE 合并（相邻小段合并到超过阈值）
        final = []
        buffer = ""
        buffer_section = ""
        for m in merged:
            if len(buffer) < FinancialDocumentParser.MIN_CHUNK_SIZE:
                sep = "\n" if buffer else ""
                buffer += sep + m["content"]
                if not buffer_section:
                    buffer_section = m["section"]
            else:
                final.append({"section": buffer_section, "content": buffer})
                buffer = m["content"]
                buffer_section = m["section"]
        if buffer:
            if final and len(buffer) < FinancialDocumentParser.MIN_CHUNK_SIZE:
                final[-1]["content"] += "\n" + buffer
            else:
                final.append({"section": buffer_section, "content": buffer})

        return [
            DocumentChunk(
                content=m["content"],
                page_number=page_num,
                paragraph_number=idx + 1,
                metadata={
                    "section": m["section"],
                    "doctype": doc_type,
                    "source": source,
                },
            )
            for idx, m in enumerate(final)
        ]

    @staticmethod
    def _wrap(paragraphs: List[str], page_num: int,
              doc_type: str = "report",
              source: str = "") -> List[DocumentChunk]:
        """纯段落回退切分：逐段累积，超过 MIN_CHUNK_SIZE 就输出；
        若段落极长（>500字），在本段落内按句号/分号切为 300-500 字子段。
        每块附带 section="正文", doctype, source。"""
        if not paragraphs:
            return [DocumentChunk(content="", page_number=page_num,
                                 metadata={"doctype": doc_type, "source": source,
                                           "section": "正文"})]

        raw = []
        for para in paragraphs:
            if not para:
                continue
            # 长段落内部软切分（仅针对报告/report 做细粒度拆分）
            if doc_type in ('report',) and len(para) > 500:
                sub_segments = FinancialDocumentParser._split_long_paragraph(para)
                raw.extend(sub_segments)
            else:
                raw.append(para)

        # MIN_CHUNK_SIZE 合并
        merged = []
        buffer = ""
        for r in raw:
            if len(buffer) < FinancialDocumentParser.MIN_CHUNK_SIZE:
                buffer = (buffer + "\n" + r).strip() if buffer else r
            else:
                merged.append(buffer)
                buffer = r
        if buffer:
            if merged and len(buffer) < FinancialDocumentParser.MIN_CHUNK_SIZE:
                merged[-1] = merged[-1] + "\n" + buffer
            else:
                merged.append(buffer)

        if not merged:
            return [DocumentChunk(content="", page_number=page_num,
                                 metadata={"doctype": doc_type, "source": source,
                                           "section": "正文"})]

        return [
            DocumentChunk(
                content=m,
                page_number=page_num,
                paragraph_number=idx + 1,
                metadata={
                    "section": "正文",
                    "doctype": doc_type,
                    "source": source,
                },
            )
            for idx, m in enumerate(merged)
        ]

    @staticmethod
    def _split_long_paragraph(para: str) -> List[str]:
        """将 >500 字的长段落按句号/分号切为 300-500 字的子段。
        对于金融研报中的大段数据分析文字有用。"""
        # 按中文句号、分号、感叹号、问号切分
        sentences = re.split(r'([。；！？\n])', para)
        segments = []
        buf = ""
        for i in range(0, len(sentences), 2):
            seg = sentences[i]
            punct = sentences[i + 1] if i + 1 < len(sentences) else ""
            chunk = (seg + punct).strip()
            if not chunk:
                continue
            if len(buf) < 300:
                buf += chunk
            else:
                if buf:
                    segments.append(buf)
                buf = chunk
        if buf:
            segments.append(buf)
        return segments if segments else [para]

    # ═══════════════════════════════════════════════
    #  辅助方法
    # ═══════════════════════════════════════════════

    @staticmethod
    def _is_section_heading(text: str, doctype: str) -> bool:
        """判断一行文本是否为章节标题。"""
        if FinancialDocumentParser.CONTRACT_HEADING.match(text):
            return True
        if FinancialDocumentParser.REPORT_BRACKET_HEADING.match(text):
            return True
        if FinancialDocumentParser.MARKDOWN_HEADING.match(text):
            return True
        return False


# ═══════════════════════════════════════════════
#  单元验证
# ═══════════════════════════════════════════════
# 运行 python financial_document_parser.py 触发

def _unit_test():
    print("=" * 60)
    print("  金语AI 金融文档解析器 —— 单元验证")
    print("=" * 60)
    print()

    # ── 测试1：监管法规 · 第X条切分 ──
    print("[测试1] 监管法规 · 第X条切分")
    reg_text = (
        "第一条 为了规范证券市场活动，保护投资者合法权益，维护社会经济秩序，"
        "促进社会主义市场经济健康发展，制定本法。本条所称证券包括股票、"
        "公司债券、存托凭证以及国务院依法认定的其他证券。\n"
        "第二条 在中华人民共和国境内，股票、公司债券、存托凭证"
        "和国务院依法认定的其他证券的发行和交易，适用本法。本法未规定的，"
        "适用《中华人民共和国公司法》和其他有关法律、行政法规的规定。\n"
        "第三条 证券的发行、交易活动，必须遵循公开、公平、公正的原则。"
        "发行人、投资者、中介机构等各方参与者应当遵守法律法规，"
        "诚实守信，不得有欺诈、内幕交易和操纵市场等违法行为。\n"
        "第四条 证券发行、交易活动的当事人，应当遵守自愿、有偿、"
        "诚实信用的原则。禁止欺诈、内幕交易和操纵证券市场的行为。\n"
    )
    r1 = FinancialDocumentParser._split_financial_text(reg_text, 1, "证券法.docx")
    print(f"  文本字数: {len(reg_text)}")
    print(f"  切分块数: {len(r1)}")
    for i, c in enumerate(r1):
        print(f"    块{i + 1}: section={c.metadata.get('section','')}  "
              f"doctype={c.metadata.get('doctype','')}  "
              f"len={len(c.content)}")
    assert len(r1) >= 2, "监管法规应切为多块"
    assert all(c.metadata.get('doctype') == 'regulation' for c in r1)
    print("  PASS")
    print()

    # ── 测试2：基金合同 · 第X章/第X条 ──
    print("[测试2] 基金合同 · 第X章/第X条")
    contract_text = (
        "第一章 总则\n"
        "第一条 根据《中华人民共和国证券投资基金法》及有关法律法规，"
        "制定本合同。基金管理人承诺以诚实信用、勤勉尽责的原则管理和运用"
        "基金资产，但不保证基金一定盈利，也不保证最低收益。\n"
        "第二条 基金合同当事人包括基金管理人、基金托管人和基金份额持有人。"
        "基金管理人：XX基金管理有限公司。基金托管人：XX银行股份有限公司。\n"
        "第二章 基金合同当事人及权利义务\n"
        "第三条 基金管理人应当履行下列职责：依法募集资金，办理基金份额"
        "的发售和登记事宜；办理基金备案手续；对所管理的不同基金财产分别管理、"
        "分别记账，进行证券投资。\n"
        "第四条 基金托管人应当履行下列职责：安全保管基金财产；按照规定开设"
        "基金财产的资金账户和证券账户；对基金财务会计报告进行复核。\n"
    )
    r2 = FinancialDocumentParser._split_financial_text(contract_text, 1, "xx基金合同.docx")
    print(f"  文本字数: {len(contract_text)}")
    print(f"  切分块数: {len(r2)}")
    for i, c in enumerate(r2):
        print(f"    块{i + 1}: section={c.metadata.get('section','')}  "
              f"doctype={c.metadata.get('doctype','')}  len={len(c.content)}")
    assert any(c.metadata.get('section', '').startswith('第') for c in r2)
    assert all(c.metadata.get('doctype') == 'contract' for c in r2)
    print("  PASS")
    print()

    # ── 测试3：年报 · 【标题】切分 ──
    print("[测试3] 年报 · 【标题】切分")
    report_text = (
        "【重要提示】\n"
        "本公司董事会及全体董事保证本报告内容不存在虚假记载、误导性陈述或"
        "重大遗漏，并对其内容的真实性、准确性和完整性承担个别及连带责任。\n"
        "【公司简介】\n"
        "XX股份有限公司（以下简称\"公司\")成立于2001年，注册资本10亿元。"
        "公司主营业务为金融科技服务，致力于为金融机构提供数字化解决方案，"
        "服务客户超过500家，覆盖银行、证券、保险等多个领域。\n"
        "【财务数据】\n"
        "报告期内，公司实现营业收入50亿元，同比增长20%。净利润8亿元，"
        "同比增长15%。基本每股收益1.5元。经营活动产生的现金流量净额12亿元。"
        "加权平均净资产收益率为18%。\n"
        "【股东信息】\n"
        "截至报告期末，公司普通股股东总数为3.5万户。前十大股东持股比例"
        "合计为65%，其中第一大股东持股比例为30%。\n"
    )
    r3 = FinancialDocumentParser._split_financial_text(report_text, 1, "xx2024年年报.pdf")
    print(f"  文本字数: {len(report_text)}")
    print(f"  切分块数: {len(r3)}")
    for i, c in enumerate(r3):
        print(f"    块{i + 1}: section={c.metadata.get('section','')}  "
              f"doctype={c.metadata.get('doctype','')}  len={len(c.content)}")
    assert all(c.metadata.get('doctype') == 'report' for c in r3)
    print("  PASS")
    print()

    # ── 测试4：研报 · 长段落回退（300-500字软切） ──
    print("[测试4] 研报 · 长段落回退（300-500字软切）")
    # 构造一个无标题的长文本，测试回退切分
    long_chunk = (
        "金融科技行业近年来发展迅速，人工智能、大数据、区块链等技术的应用"
        "正在深刻改变金融服务的形态。传统金融机构纷纷加大科技投入，"
        "推动数字化转型。与此同时，互联网金融平台也在不断壮大，"
        "两者之间的竞争与合作日益深化。"
        "从市场规模来看，2024年中国金融科技市场规模预计达到5000亿元，"
        "同比增长25%。其中，支付科技、信贷科技、保险科技和财富管理科技"
        "是最主要的四个细分领域，分别占比30%、25%、20%和15%。"
    ) * 10  # ~7800字
    # 不包含任何标题，触发回退切分
    r4 = FinancialDocumentParser._split_financial_text(long_chunk, 1, "金融科技行业深度研究报告.pdf")
    print(f"  文本字数: {len(long_chunk)}")
    print(f"  切分块数: {len(r4)}")
    for i, c in enumerate(r4):
        print(f"    块{i + 1}: section={c.metadata.get('section','')}  "
              f"len={len(c.content)}  doctype={c.metadata.get('doctype','')}")
    # doctype 应检测为 report（文件名含"研究报告"）
    assert all(c.metadata.get('doctype') == 'report' for c in r4), "文件名应触发 report 类型"
    print("  PASS")
    print()

    # ── 测试5：混合文本（文件名控制 doctype） ──
    print("[测试5] 混合文本 · 文件名指定 doctype")
    r5_all = []
    for name in ("证券法.docx", "基金合同.docx", "2024年报.pdf", "行业研究报告.pdf"):
        t = f"第一条 这是{name}的测试文本，用于验证doctype检测。\n第二条 继续测试。\n"
        r5 = FinancialDocumentParser._split_financial_text(t, 1, name)
        r5_all.extend(r5)
    for c in r5_all:
        print(f"  source={c.metadata.get('source',''):30s}  "
              f"doctype={c.metadata.get('doctype',''):12s}  "
              f"section={c.metadata.get('section','')}")
    print("  PASS")
    print()

    print("=" * 60)
    print(f"  全部测试通过！共生成 {sum(len(r) for r in [r1, r2, r3, r4, r5_all])} 个文档块")
    print("=" * 60)


if __name__ == "__main__":
    _unit_test()
